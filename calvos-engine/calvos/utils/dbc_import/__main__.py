#!/usr/local/bin/python2.7
# encoding: utf-8
'''
CalvOS - Open Source SW Utilitites for Embedded Systems.

Calvos provides a set of SW utilitites for the development
of embedded systems in C-language.

@author:     Carlos Calvillo

@copyright:  2020 Carlos Calvillo. All rights reserved.

@license:    GPL v3

@contact:    calcore.io@gmail.com
@deffield    updated: Updated
'''           
import sys
import os
import pathlib as pl
import shutil
import logging
import traceback
import csv

from argparse import ArgumentParser
from argparse import RawDescriptionHelpFormatter

__all__ = []
__date__ = '2021-10-26'
__updated__ = '2020-10-26'

DEBUG = 1
TESTRUN = 0
PROFILE = 0

class CLIError(Exception):
    '''Generic exception to raise and log different fatal errors.'''
    def __init__(self, msg):
        super(CLIError).__init__(type(self))
        self.msg = "E: %s" % msg
    def __str__(self):
        return self.msg
    def __unicode__(self):
        return self.msg
    

class CAN_Signal():
    def __init__(self, name = None, def_val = None, fail_val = None, msg_name = None):
        self.name = name
        self.def_val = def_val
        self.fail_val = fail_val
        self.msg_name = msg_name

class SigSection():
    
    def __init__(self, signals = [], paragraphs = []):
        
        self.signals = signals # List of CAN_Signal objects
        self.paragraphs = [] # List of docx Paragraph objects
        self.output = [] # List of docx Paragraph objects
        
    def gen_output(self):
        
        self.output = []
        
        for signal in self.signals:
            for paragraph in self.paragraphs:
                temp_str = paragraph.text
                if signal.name is not None:
                    temp_str.replace("${signal.name}", signal.name)
                if signal.def_val is not None:
                    temp_str.replace("${signal.def_val}", signal.def_val)
                if signal.fail_val is not None:
                    temp_str.replace("${signal.fail_val}", signal.fail_val)
                if signal.msg_name is not None:
                    temp_str.replace("${message.name}", signal.msg_name)
                
                paragraph.text = temp_str
                
                self.output.append(paragraph)
    
    def get_output(self):
        
        return self.output
        
        
class MsgSection():
    
    def __init__(self, start_idx = None, end_idx = None):
        
        self.start_idx = start_idx
        self.end_idx = end_idx
        
        signal_sections = [] 
        


def string_to_path(path_string):
    ''' Returns a "path" object for the given string. '''
    # pathlib requires regular diagonal, not reverse diagonal
    path_string = path_string.replace("\\", "/")
    
    return pl.Path(path_string)
    
    
def folder_exists(path):
    ''' Returns "True" if the given path is an existing folder. '''
    return_value = False
    
    if type(path) is str:
        path = string_to_path(path)
     
    if path.exists() and path.is_dir():
        return_value = True
    
    return return_value


def file_exists(path):
    ''' Returns "True" if the given path is an existing file. '''
    return_value = False
    
    if type(path) is str:
        path = string_to_path(path)
     
    if path.exists() and path.is_file():
        return_value = True
    
    return return_value

def create_folder(folder):
    """ Creates the specified folder.
    """
    if folder_exists(folder) is False:
        try:
            folder.mkdir()
        except Exception as e:
            print('Failed to create folder %s. Reason: %s' % (folder, e))
    else:
        print("Folder to be created '%s' already exists." % folder)
    

def main(argv=None): # IGNORE:C0111
    '''Command line options.'''
    # Import version from __version__.py file
    # Set calvos path which will be used for accessing package resources
    calvos_path = pl.Path(__file__).parent.absolute()
    version_path = calvos_path / "__version__.py"
        
    about = {}
    with open(version_path) as f:
        exec(f.read(), about)
    __version__ = about["__version__"]

    if argv is None:
        argv = sys.argv
    else:
        sys.argv.extend(argv)

    program_name = os.path.basename(sys.argv[0])
    program_version = "v%s" % __version__
    program_build_date = str(__updated__)
    program_version_message = 'calvos %s' % (program_version)
    program_shortdesc = __import__('__main__').__doc__.split("\n")[1]
    program_license = '''%s
Copyright (C) 2020  Carlos Calvillo
  
  This program is free software: you can redistribute it and/or modify
  it under the terms of the GNU General Public License as published by
  the Free Software Foundation, either version 3 of the License, or
  (at your option) any later version.
 
  This program is distributed in the hope that it will be useful,
  but WITHOUT ANY WARRANTY; without even the implied warranty of
  MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
  GNU General Public License for more details.  

  You should have received a copy of the GNU General Public License
  along with this program.  If not, see <https://www.gnu.org/licenses/>.

USAGE
''' % program_shortdesc

    try:
        
        # Check if calvos package is installed by trying to import the logging system
        try:
            import calvos.common.logsys as lg
        except:
            print("INFO: Setting up package...")
            # Add package path to sys.path and attempt the importing again
            calvos_pkg_path_str = str(calvos_path.parent)
            if calvos_pkg_path_str not in sys.path:
                sys.path.append(calvos_pkg_path_str)
                print("INFO: calvos path added to sys.path.")
            # Attempt the import
            try:
                import calvos.common.logsys as lg
            except Exception as e:
                print("ERROR: Couldn't load calvos package.")
                print("Generated exception:\n")
                print(str(e))
                return 2
        

        log_level = logging.INFO
        
        
        
        # Setup argument parser
        parser = ArgumentParser(description=program_license, \
                                formatter_class=RawDescriptionHelpFormatter)
        parser.add_argument("-i","--input", dest="dbc_in", required=False, \
            help=("DBC file to process"))
        parser.add_argument("-t","--template", dest="template_in", required=False, \
            help=("DBC file to process"))
        parser.add_argument("-o","--out", dest="doc_out", required=False, \
            help=("DBC file to process"))
        
        
#         parser.add_argument('-V', '--version', action='version', version=program_version_message)
#         parser.add_argument('-v', '--ver', action='version', version=program_version_message)
#         parser.add_argument("-p","--project", dest="project", \
#             required=(('--demo' not in sys.argv) and ('-d' not in sys.argv)), \
#             help=("Required (if -d was not provided). Full path with file name of the calvos" \
#                   + "project to be processed."))
#         parser.add_argument("-l","--log", dest="log_level", required=False, \
#             help=("Optional. LOG_LEVEL: 0 - Debug, 1 - Info, "\
#                   + "2 - Warning, 3 - Error. Default is 1 - Info."))
#         parser.add_argument("-e","--export", dest="export", required=False, \
#             help="Optional. Generated C-code will be exported (copied) into the provided EXPORT path.")
#         parser.add_argument("-b","--backup", dest="backup", required=False, \
#             help=("Optional. Backups of the overwritten C-code during an export operation " \
#                   + "will be placed in the provided BACKUP path. " \
#                   + "This is only used if -e argument was provided."))
#         parser.add_argument("--dont-export-user", dest="dont_export_user", action='store_true', \
#                             required=False, \
#             help=("Optional. If argument -e is provided and also this one, then the export " \
#                   + "operation won't export 'USER_' files. Use with care."))
        

        # Process arguments
        args = parser.parse_args()
        dbc_in = args.dbc_in
        template_in = args.template_in
        doc_out = args.doc_out
        
        arguments_OK = True
        
        
        dbc_in = string_to_path(dbc_in)
        
        if file_exists(dbc_in):
            
            demo_path = dbc_in.parent.resolve()
            
            print(demo_path)
            # Setting up logging system
            log_output_file = demo_path / "log.log"
            if file_exists(log_output_file):
                log_output_file.unlink()
            lg.log_system = lg.Log(log_level, log_output_file)
            log = lg.log_system
            log.add_logger("main")
            
            import calvos.comgen.CAN as CAN
            import calvos.common.project as pj
            import canmatrix
            
            proj = pj.Project("Project Name", None, calvos_path)
            
            can_object = CAN.Network_CAN(proj)
        
            can_dbc = canmatrix.CanMatrix()
            
            can_dbc = canmatrix.formats.loadp(str(dbc_in))
            
            
            msg_sections = []
            
            
            if file_exists(template_in):
                import docx
                
                doc = docx.Document(template_in)
                document = docx.Document()
                
                
                open_msg_section = False
                open_sig_section = False
                section_idxs = []
                
                # Identify message sections
                
                for idx, par in enumerate(doc.paragraphs):
                    if "<<<for-rx-message>>>" in par.text and open_msg_section is False:
                        open_msg_section = True
                        section_idxs.append(idx)
                    elif "<<<for-rx-message>>>" in par.text and open_msg_section is True:
                        print("ERROR: nesting <<<for-rx-message>>> sections not allowed.")
                        return 0
                    elif "<<<end-for-rx-message>>>" in par.text and open_msg_section is True:
                        open_msg_section = False
                        section_idxs.append(idx)
                        # Create msg_section
                        msg_sections.append(MsgSection(section_idxs[0],section_idxs[1]))
                    elif "<<<end-for-rx-message>>>" in par.text and open_msg_section is False:
                        print("ERROR: Found <<<end-for-rx-message>>> without opening counter part")
                        return 0
                    else:
                        pass
                    
                    
                # identify signal sections
                for msg_section in msg_sections:
                    for paragraph in msg_section
                    
                
                
                document.add_paragraph(par.text, par.style)
                
                
                
                
                
                
                
                
                
                document.save(doc_out)
                
                
                for idx, paragraph in enumerate(doc.paragraphs):
                    
                    doc.paragraphs[idx].text =  doc.paragraphs[idx].text + "hhh"
                    print(idx, "   ", paragraph.text)

            
#             for dbc in can_dbc:
#                 for msg in can_dbc[dbc].frames:
#                     if "BCM" in msg.receivers:
#                         print(msg.name)
#                         for signal in msg.signals:
#                             if "BCM" in signal.receivers:
#                                 print("    ", signal.name)

            
            print(dbc_in)
            
            
        else:
            print(" Input file not found.")
        
        print("Hello")
   
        return 0
    except KeyboardInterrupt:
        ### handle keyboard interrupt ###
        return 0
    except Exception as e:
        if DEBUG or TESTRUN:
            raise(e)
        indent = len(program_name) * " "
        
        message = program_name + ": " + repr(e) + "\n"
        message = message + traceback.format_exc()
        sys.stderr.write(program_name + ": " + repr(e) + "\n")
        sys.stderr.write(indent + "  for help use --help")
        
        return 2

if __name__ == "__main__":
    if DEBUG:
        
        sys.argv.append("-i")
        sys.argv.append("G:\\devproj\\github\\dbc_import\\calvos\\calvos-engine\\calvos\\demo\\test.dbc")
        
        sys.argv.append("-t")
        sys.argv.append("G:\\devproj\\github\\dbc_import\\calvos\\calvos-engine\\calvos\\demo\\Template.docx")
        
        sys.argv.append("-o")
        sys.argv.append("G:\\devproj\\github\\dbc_import\\calvos\\calvos-engine\\calvos\\demo\\Template_out.docx")
        
        pass
    if TESTRUN:
        import doctest
        doctest.testmod()
    if PROFILE:
        import cProfile
        import pstats
        profile_filename = 'calvos.main_profile.txt'
        cProfile.run('main()', profile_filename)
        statsfile = open("profile_stats.txt", "wb")
        p = pstats.Stats(profile_filename, stream=statsfile)
        stats = p.strip_dirs().sort_stats('cumulative')
        stats.print_stats()
        statsfile.close()
        sys.exit(0)
    sys.exit(main())